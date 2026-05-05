"""Generate LINE-2D Shape-Based Matching technical document as .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 51, 102)

# Code style
if 'Code' not in [s.name for s in doc.styles]:
    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)

def add_code(text):
    p = doc.add_paragraph(text, style='Code')
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for row in rows:
        r = t.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = str(val)
            for p in r.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))

def add_numbered(text):
    doc.add_paragraph(text, style='List Number')

# ===== TITLE PAGE =====
doc.add_paragraph('')
doc.add_paragraph('')
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('LINE-2D Shape-Based Matching')
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0, 51, 102)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Technical Handover Document')
r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(80, 80, 80)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('Wafer Alignment System — QES (Asia-Pacific) Sdn Bhd')
r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(100, 100, 100)

t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run('May 2026')
r4.font.size = Pt(12); r4.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ===== 1. EXECUTIVE SUMMARY =====
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This document describes the LINE-2D shape-based template matching subsystem used in the '
    'Wafer Alignment application. The algorithm locates a trained template pattern in a search '
    'image by comparing quantised gradient orientations. It is a Python port of the C++ SBM_line2D '
    'implementation (reference: SeanYan604/SBM_line2D).')
doc.add_paragraph(
    'Key capabilities: rotation-invariant matching (0°–360°, configurable step), multi-scale support, '
    'sub-pixel position and angle refinement, and coarse-to-fine pyramid acceleration for '
    'high-resolution (>2000 px) images.')

# ===== 2. ARCHITECTURE =====
doc.add_heading('2. Architecture Overview', level=1)
doc.add_heading('2.1 MVVM Layer Mapping', level=2)
add_table(['Layer', 'File', 'Role'], [
    ['Service', 'linemod_matcher.py', 'Core algorithm (1330 lines)'],
    ['ViewModel', 'pattern_viewmodel.py', 'UI logic, config, caching'],
    ['View', 'pattern_tab.py', 'Tkinter widgets'],
    ['Model', 'recipe_model.py', 'XML recipe persistence'],
    ['Model', 'app_state.py', 'Shared mutable state'],
    ['Integration', 'zmq_server.py', 'ZeroMQ TCP interface'],
])

doc.add_heading('2.2 Dependency Flow', level=2)
doc.add_paragraph(
    'pattern_tab (View) → PatternViewModel → LinemodMatcher (Service)\n'
    'PatternViewModel → RecipeManager (Model)\n'
    'ZMQ Server → LinemodMatcher + RecipeManager')

# ===== 3. ALGORITHM PIPELINE =====
doc.add_heading('3. Algorithm Pipeline', level=1)
doc.add_paragraph(
    'The matching pipeline has two phases: offline (template generation) and online '
    '(search-image matching).')

doc.add_heading('3.1 Phase 1 — Template Generation (generate_templates)', level=2)
items = [
    'Rotation & Scale Grid: For each (angle, scale) pair the base template image is rotated and '
    'scaled via cv2.warpAffine. Default: 5° step × 1 scale = 72 templates.',
    'Mask Handling: A rotation mask (all-white image transformed identically) prevents border '
    'artefacts. An optional user-drawn detection mask restricts feature extraction.',
    'Pyramid Construction: For each pyramid level l ∈ {0,1,2}, the image is downsampled by 2^l '
    'via cv2.pyrDown.',
    'Gradient Quantisation: At each level, Sobel gradients → angle → 8 orientation bins.',
    'Feature Extraction: N spatially-scattered features selected by gradient magnitude.',
    'Bounding-Box Crop: _crop_templates computes the tightest bounding box and makes feature '
    'coordinates relative.',
    'NumPy Pre-caching: Feature (x,y,label) arrays cached as int32 vectors for inner-loop speed.',
]
for item in items:
    add_numbered(item)

doc.add_paragraph(
    'Template generation is parallelised with ThreadPoolExecutor (up to 8 workers). '
    'OpenCV/NumPy release the GIL so threads run in parallel.')

doc.add_heading('3.2 Phase 2 — Online Matching (match)', level=2)
doc.add_paragraph('Two execution paths exist, selected automatically by image size:')
add_bullet('Single-Level (≤ 2000 px): Full-resolution scan.')
add_bullet('Pyramid (> 2000 px): Coarse-to-fine search.')

# ===== 4. CORE ALGORITHMS =====
doc.add_heading('4. Core Algorithms', level=1)

doc.add_heading('4.1 Gradient Quantisation', level=2)
doc.add_paragraph('Function: _quantize_gradients(src, weak_threshold, fast_mode, kernel_size)')
items = [
    'Convert to float32 grayscale.',
    'Optional Gaussian blur (auto-sigma for images >1500 px, or explicit kernel).',
    'Sobel derivatives: dx = Sobel(I, x), dy = Sobel(I, y), ksize=3.',
    'Magnitude: M = sqrt(dx² + dy²); Angle: θ = atan2(dy, dx) mod 360°.',
    'Quantise to 16 bins then fold to 8: quant_8 = (angle × 16/360).astype(uint8) & 7.',
    'Adaptive threshold: If weak_threshold < 0, use the |value|-th percentile of non-zero magnitudes.',
]
for item in items:
    add_numbered(item)

doc.add_heading('Hysteresis Mode (fast_mode=False)', level=3)
doc.add_paragraph(
    'For each of 8 orientation bins, count votes in a ks×ks neighbourhood via cv2.filter2D. '
    'Accept a pixel only if M > threshold AND the winning bin has ≥ ⌈ks²/2⌉ + 1 votes. '
    'Kernel size: auto = min(max(3, ⌊3 × maxdim/1500⌋) | 1, 9).')

doc.add_heading('Fast Mode (fast_mode=True)', level=3)
doc.add_paragraph(
    'Skip hysteresis; accept any pixel where M > threshold. Used for search images where speed matters.')

doc.add_paragraph('Output: A uint8 image where each pixel is a power-of-2 bit (2^bin) or 0.')

doc.add_heading('4.2 Orientation Spreading', level=2)
doc.add_paragraph('Function: _spread(quantized, T)')
doc.add_paragraph(
    'OR-dilates each of the 8 bit-planes independently with a (2T−1) × (2T−1) rectangular kernel '
    'via cv2.dilate, then ORs the results. This creates tolerance to small spatial misalignment.')
doc.add_paragraph('Default T values per pyramid level: [4, 8, 16].')

doc.add_heading('4.3 Response Map LUTs', level=2)
doc.add_paragraph('Function: _build_response_luts() → 8 tables of 256 entries.')
doc.add_paragraph(
    'For each template label L ∈ {0..7} and each possible spread byte v ∈ {0..255}:')
add_bullet('LUT[L][v] = 4  if bit L is set in v (exact match)')
add_bullet('LUT[L][v] = 1  if a neighbour bit (L±1 mod 8) is set (neighbour match)')
add_bullet('LUT[L][v] = 0  otherwise')
doc.add_paragraph(
    'Response maps are computed via cv2.LUT(spread_img, LUT[label]) — a single vectorised lookup per label.')

doc.add_heading('4.4 Scattered Feature Selection', level=2)
doc.add_paragraph('Function: _extract_scattered_features(quantized, magnitude, N, mask)')
items = [
    'Collect all non-zero quantised pixels (optionally masked).',
    'Sort by descending gradient magnitude.',
    'Greedily select N features with minimum-distance constraint. Distance starts at '
    'd = perimeter / (0.8N) and shrinks by 15% per pass.',
    'Uses spatial grid hashing: distance checks are O(9 × k) per candidate instead of O(n).',
]
for item in items:
    add_numbered(item)

doc.add_heading('4.5 Score Map Computation', level=2)
doc.add_paragraph('For each template, accumulate response values at feature locations:')
add_code('S(r,c) = Σ ResponseMap[label_i][r + y_i, c + x_i]  for i=1..N_feat')
add_code('similarity(r,c) = 100 × S(r,c) / (4 × N_valid)')
doc.add_paragraph(
    'Features are sorted by label to improve L2 cache hit rate (same response map accessed consecutively). '
    'Score maps use int16 dtype (max 128 × 4 = 512 < 32767).')

doc.add_heading('4.6 Sub-Pixel Position Refinement', level=2)
doc.add_paragraph('Function: _subpixel_refine(score_map, r, c)')
doc.add_paragraph('3-point parabolic fit on each axis independently:')
add_code('δx = (S[r,c-1] - S[r,c+1]) / (2 × (S[r,c-1] - 2×S[r,c] + S[r,c+1]))')
doc.add_paragraph('Offset clamped to [-0.5, +0.5].')

doc.add_heading('4.7 Angular Sub-Pixel Interpolation', level=2)
doc.add_paragraph('Function: _angular_interpolate(matches)')
doc.add_paragraph(
    'Given the best discrete match at angle θ₀ with score S₀, and scores S₋, S₊ from the '
    'adjacent angle templates:')
add_code('δθ = (Δθ/2) × (S₋ - S₊) / (S₋ - 2×S₀ + S₊),  |δθ| ≤ Δθ/2')

doc.add_heading('4.8 Non-Maximum Suppression', level=2)
doc.add_paragraph('Function: _nms(matches)')
doc.add_paragraph(
    'Sort by score descending. Keep a match only if its Euclidean distance to all previously '
    'kept matches exceeds NMS_DISTANCE (default 30 px).')

# ===== 5. PYRAMID MATCHING =====
doc.add_heading('5. Pyramid Matching', level=1)
doc.add_paragraph('Function: _match_pyramid(search_gray, threshold)')

doc.add_heading('5.1 Coarse Pass (Level N)', level=2)
items = [
    'Downsample search image by 2^N via cv2.pyrDown. Level N = min(PYRAMID_LEVELS−1, ⌈log₂(maxdim/1200)⌉).',
    'Quantise, spread (with T_N), compute response maps at coarse resolution.',
    'Score each template using only the first COARSE_NUM_FEATURES (default 32) features.',
    'Lower threshold: max(0.3 × threshold, 15).',
    'Keep top candidate per template. Cap total to MAX_FINE_CANDIDATES.',
]
for item in items:
    add_numbered(item)

doc.add_heading('5.2 Fine Pass (Level 0)', level=2)
items = [
    'Cluster overlapping ROI windows into super-ROIs.',
    'For each super-ROI: quantise (fast mode), spread, build response maps once.',
    'Score each candidate template within its sub-window of the super-ROI.',
    'Full feature set (N=128) at full resolution for precise localisation.',
]
for item in items:
    add_numbered(item)

# ===== 6. CONFIG REFERENCE =====
doc.add_heading('6. Configuration Reference', level=1)
add_table(['Parameter', 'Default', 'Description'], [
    ['ANGLE_STEP', '5', 'Degrees between rotation templates'],
    ['SCALE_MIN/MAX', '1.0/1.0', 'Scale range (0.8–1.2 for Full Search)'],
    ['SCALE_STEP', '0.1', 'Scale increment'],
    ['WEAK_THRESHOLD', '-70', 'Gradient threshold. Negative = percentile-adaptive'],
    ['NUM_FEATURES', '128', 'Features per template level'],
    ['HYSTERESIS_KERNEL', '0', 'Majority-vote kernel size (0=auto)'],
    ['T_PYRAMID', '[4,8,16]', 'Spread T per pyramid level'],
    ['MATCH_THRESHOLD', '50.0', 'Minimum similarity (0–100)'],
    ['NMS_DISTANCE', '30', 'Non-max suppression radius (px)'],
    ['FAST_SEARCH_QUANTIZE', 'True', 'Skip hysteresis on search images'],
    ['COARSE_NUM_FEATURES', '32', 'Features for coarse pass'],
    ['MAX_COARSE_CANDIDATES', '2', 'Top candidates per template (coarse)'],
    ['MAX_FINE_CANDIDATES', '5', 'Cap before fine pass'],
])

# ===== 7. SEARCH MODES =====
doc.add_heading('7. Search Modes (UI)', level=1)
add_table(['Mode', 'ANGLE_STEP', 'Scale', 'Templates'], [
    ['Simple (Fast)', '360 (only 0°)', '1.0', '1'],
    ['With Rotation', '10', '1.0', '36'],
    ['Full Search', '10', '0.8–1.2', '~180'],
])

# ===== 8. DATA STRUCTURES =====
doc.add_heading('8. Data Structures', level=1)

doc.add_heading('8.1 Feature', level=2)
add_code("class Feature:\n    __slots__ = ['x', 'y', 'label']  # x,y: position; label: orientation bin 0-7")

doc.add_heading('8.2 TemplatePyr', level=2)
add_code("class TemplatePyr:\n    __slots__ = ['width', 'height', 'tl_x', 'tl_y',\n"
         "                 'pyramid_level', 'features',\n"
         "                 'feat_xs_arr', 'feat_ys_arr', 'feat_labels_arr']")

doc.add_heading('8.3 Template Pyramid Entry (dict)', level=2)
add_code("{ 'angle': float, 'scale': float,\n"
         "  'templates': [TemplatePyr, ...],  # one per pyramid level\n"
         "  'size': (width, height) }")

doc.add_heading('8.4 Match Result (dict)', level=2)
add_code("{ 'x': float, 'y': float,          # sub-pixel centre\n"
         "  'angle': float,                   # refined angle (degrees)\n"
         "  'angle_discrete': float,          # original grid angle\n"
         "  'angle_interp_delta': float,      # interpolation offset\n"
         "  'scale': float, 'score': float,   # 0-100 similarity\n"
         "  'template_id': int,\n"
         "  'bbox': (x, y, w, h) }")

# ===== 9. ZMQ INTEGRATION =====
doc.add_heading('9. ZMQ Integration', level=1)

doc.add_heading('9.1 PM_REQ Command', level=2)
add_code('PM_REQ "<image_path>" "<recipe_path>"')
doc.add_paragraph('Response (JSON):')
add_code('{"status":"ok", "x":..., "y":..., "angle":..., "score":..., "delta_x":..., "delta_y":...}')

doc.add_heading('9.2 TEACH_REQ Command', level=2)
add_code('TEACH_REQ "<image_path>" "<recipe_path>"')
doc.add_paragraph(
    'Opens interactive OpenCV ROI selector. Operator crops template, optionally draws detection mask. '
    'Saves to recipe XML.')

# ===== 10. RECIPE XML =====
doc.add_heading('10. Recipe XML Schema (FindPattern)', level=1)
add_code('<FindPattern>\n'
         '  <MatchThreshold>50.0</MatchThreshold>\n'
         '  <NumFeatures>128</NumFeatures>\n'
         '  <GradThrPct>70.0</GradThrPct>\n'
         '  <TSpread>4</TSpread>\n'
         '  <HystKernel>0</HystKernel>\n'
         '  <SearchMode>Simple (Fast)</SearchMode>\n'
         '  <TemplatePath>C:\\...\\template.png</TemplatePath>\n'
         '  <TemplateCropCX>512.0</TemplateCropCX>\n'
         '  <TemplateCropCY>310.0</TemplateCropCY>\n'
         '  <DetectionMaskPath></DetectionMaskPath>\n'
         '</FindPattern>')

# ===== 11. PERFORMANCE =====
doc.add_heading('11. Performance Characteristics', level=1)
add_table(['Phase', '919px', '2560px', '5120px'], [
    ['Quantise (fast)', '5 ms', '15 ms', '50 ms'],
    ['Spread (T=4)', '8 ms', '25 ms', '90 ms'],
    ['Response Maps', '3 ms', '10 ms', '35 ms'],
    ['Scoring (1 tmpl)', '10 ms', '80 ms', '300 ms'],
    ['Scoring (72 tmpl)', '200 ms', '—', '—'],
    ['Pyramid total', '—', '180 ms', '400 ms'],
])
doc.add_paragraph(
    'Memory: Response maps use int16 (halves bandwidth vs int32). Rotated template images are NOT '
    'stored; regenerated on-demand for visualisation.')

# ===== 12. DETECTION MASK =====
doc.add_heading('12. Detection Mask System', level=1)
doc.add_paragraph(
    'Users can draw a polygon mask on the template to exclude regions (e.g. wafer boundary curvature).')
items = [
    'mask_editor.py: OpenCV interactive polygon drawing window with auto-detection of wafer boundary.',
    'Mask stored as <template>_mask.png alongside the template.',
    'Passed to _extract_scattered_features as a binary mask: features only extracted where mask > 0.',
    'Mask is rotated/scaled identically to the template during generate_templates.',
]
for item in items:
    add_numbered(item)

# ===== 13. CACHING =====
doc.add_heading('13. Smart Template Caching', level=1)
doc.add_paragraph('The PatternViewModel tracks a config fingerprint string:')
add_code('f"{NUM_FEATURES}_{WEAK_THRESHOLD}_{T_PYRAMID}_{HYSTERESIS_KERNEL}_'
         '{ANGLE_STEP}_{SCALE_MIN}_{SCALE_MAX}_{img_hash}_{mask_hash}"')
doc.add_paragraph(
    'Templates are only regenerated when this string changes. This avoids redundant ~2 s rebuilds '
    'during slider tuning.')

# ===== 14. LIMITATIONS =====
doc.add_heading('14. Known Limitations & Future Work', level=1)
items = [
    'No GPU acceleration (CUDA Sobel/dilate was prototyped but PCIe overhead exceeded gains for '
    'template-sized crops).',
    'Scale search is coarse (0.1 step); finer steps increase template count linearly.',
    'Single-level fallback (_match_single_level) is slow for images >2000 px; pyramid path should '
    'always be preferred.',
    'No rotation-invariant NMS (matches at different angles near the same position may both survive).',
]
for item in items:
    add_bullet(item)

# ===== 15. FILE REFERENCE =====
doc.add_heading('15. File Reference', level=1)
add_table(['File', 'Purpose'], [
    ['app/services/linemod_matcher.py', 'Core LINE-2D algorithm (1330 lines)'],
    ['app/viewmodels/pattern_viewmodel.py', 'UI business logic, config, timing chart'],
    ['app/views/pattern_tab.py', 'Tkinter tab UI (sliders, buttons, image display)'],
    ['app/views/mask_editor.py', 'Interactive polygon mask editor'],
    ['app/models/recipe_model.py', 'XML recipe read/write'],
    ['app/models/app_state.py', 'Shared state (template crop centre, mask, flags)'],
    ['app/services/zmq_server.py', 'ZMQ server (PM_REQ, TEACH_REQ handlers)'],
])

# Save
out = os.path.join(os.path.dirname(__file__), 'LINE2D_Shape_Based_Matching.docx')
doc.save(out)
print(f'Saved: {out}')
