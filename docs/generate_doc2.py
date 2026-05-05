"""Generate Gradient-Based Wafer Edge Detection technical document as .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 51, 102)

if 'Code' not in [s.name for s in doc.styles]:
    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)

def add_code(text):
    p = doc.add_paragraph(text, style='Code')
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for row in rows:
        r = t.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = str(val)
            for p in r.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))

def add_numbered(text):
    doc.add_paragraph(text, style='List Number')

# ===== TITLE PAGE =====
doc.add_paragraph('')
doc.add_paragraph('')
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Gradient-Based Wafer Edge Detection')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0, 51, 102)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Technical Handover Document')
r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(80, 80, 80)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('Wafer Alignment System — QES (Asia-Pacific) Sdn Bhd')
r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(100, 100, 100)

t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run('May 2026')
r4.font.size = Pt(12); r4.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ===== 1. EXECUTIVE SUMMARY =====
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This document describes the gradient-based wafer edge detection subsystem. The algorithm finds '
    'the precise straight-line boundary of a semiconductor wafer in microscope images, returning '
    'sub-pixel edge position and offset from the image centre. It supports four scan directions '
    '(LEFT, RIGHT, TOP, BOTTOM), three edge polarity modes, and integrates with the FOV classifier '
    'to reject non-edge images.')

# ===== 2. ARCHITECTURE =====
doc.add_heading('2. Architecture Overview', level=1)

doc.add_heading('2.1 MVVM Layer Mapping', level=2)
add_table(['Layer', 'File', 'Role'], [
    ['Service', 'edge_finder.py', 'Core algorithm (838 lines)'],
    ['Service', 'fov_classifier.py', 'FOV type classification (1418 lines)'],
    ['ViewModel', 'edge_viewmodel.py', 'UI logic, config cache'],
    ['View', 'edge_tab.py', 'Tkinter widgets'],
    ['Model', 'recipe_model.py', 'XML recipe persistence'],
    ['Model', 'app_state.py', 'Per-direction config cache'],
    ['Integration', 'zmq_server.py', 'ZeroMQ TCP interface'],
])

doc.add_heading('2.2 Dependency Flow', level=2)
doc.add_paragraph(
    'edge_tab (View) → EdgeViewModel → EdgeLineFinder (Service) → FOVClassifier\n'
    'EdgeViewModel → RecipeManager (Model)\n'
    'ZMQ Server → EdgeLineFinder + RecipeManager')

# ===== 3. PIPELINE =====
doc.add_heading('3. Algorithm Pipeline', level=1)
doc.add_paragraph('The edge detection pipeline consists of 8 sequential steps:')
steps = [
    'Image Preprocessing (resize + CLAHE)',
    'FOV Classification (edge/die/wafer)',
    'Region Partitioning (40 strips)',
    '1-D Gradient Computation (convolution kernel)',
    'Polarity Filtering (direction-aware)',
    'Edge Point Detection (cluster + sub-pixel)',
    'RANSAC Line Fitting (2000 iterations)',
    'Delta Computation (perpendicular intersection)',
]
for s in steps:
    add_numbered(s)

# ===== 4. DETAILED ALGORITHMS =====
doc.add_heading('4. Detailed Algorithm Description', level=1)

doc.add_heading('4.1 Step 1: Image Preprocessing', level=2)
doc.add_paragraph('Function: preprocess_image(img_or_path, target_dim)')
items = [
    'Load image as grayscale (cv2.imread with flag 0).',
    'Resize so the longest dimension equals TARGET_PROCESS_DIM (default 1000 px). '
    'Record scale factor s = 1000 / max(H, W).',
    'Apply CLAHE (Contrast Limited Adaptive Histogram Equalisation): clipLimit=2.0, '
    'tileGridSize=8×8.',
]
for item in items:
    add_numbered(item)
doc.add_paragraph(
    'The scale factor s is used later to convert detected coordinates back to '
    'original-resolution pixels.')

doc.add_heading('4.2 Step 2: FOV Classification', level=2)
doc.add_paragraph('Class: FOVClassifier (in fov_classifier.py)')
doc.add_paragraph(
    'Determines if the image contains a wafer edge (EDGE_FOV), die patterns (DIE_FOV), '
    'or plain wafer surface (WAFER_FOV). This step is skipped when skip_classification=True '
    '(param tuner mode or ZMQ FORCE_RUN).')
doc.add_paragraph('Classification uses four parallel detectors:')
items = [
    '1-D Edge Detection: Horizontal and vertical intensity profiles from 4 bands. '
    'Detects monotonic rising/falling regions with adaptive thresholds (percentile-based).',
    '2-D Sobel Edge Detection: Downsampled 4×, Sobel gradients, checks if strong '
    'gradients are spatially concentrated (edge) vs distributed (die).',
    'Peak-Based Classifier: Counts prominent peaks in 1-D projections. '
    '≤2 peaks → edge; ≥5 → die.',
    'Region Analysis: Splits image into 6 regions (left/centre/right + top/centre/bottom). '
    'Classifies each by texture score (Laplacian variance) and gradient peak count.',
]
for item in items:
    add_numbered(item)
doc.add_paragraph(
    'Confidence score (0–1) computed from signal strength, region consistency, and detector '
    'agreement. Below CONFIDENCE_THRESHOLD (0.40) → UNCERTAIN.')

doc.add_heading('4.3 Step 3: Region Partitioning', level=2)
doc.add_paragraph(
    'The image is divided into NUM_REGIONS (default 40) horizontal or vertical strips:')
add_bullet('LEFT/RIGHT (vertical edge): N horizontal strips, each of height H/N. '
           'Each strip produces one edge point along the x-axis.')
add_bullet('TOP/BOTTOM (horizontal edge): N vertical strips, each of width W/N. '
           'Each strip produces one edge point along the y-axis.')

doc.add_heading('4.4 Step 4: 1-D Gradient Computation', level=2)
doc.add_paragraph('For each region strip:')
items = [
    'Compute the median intensity profile along the scan axis. Median is used '
    'instead of mean for robustness against outlier pixels.',
    'Convolve with a gradient kernel: K = [-1, ..., -1, 0, 1, ..., 1] '
    'of size KERNEL_SIZE (default 7).',
]
for item in items:
    add_numbered(item)

doc.add_heading('Gradient Kernel Construction', level=3)
doc.add_paragraph('Function: create_gradient_kernel(size)')
add_code("# size=7 -> [-1, -1, -1, 0, 1, 1, 1]\n"
         "kernel = np.concatenate([-np.ones(half), np.zeros(1), np.ones(half)])")
doc.add_paragraph(
    'Multiple -1/+1 elements improve noise resistance compared to a simple [-1, 0, 1] kernel.')

doc.add_heading('4.5 Step 5: Edge Polarity Filtering', level=2)
doc.add_paragraph('Function: _apply_polarity(gradient)')
doc.add_paragraph('Three modes:')
add_table(['Mode', 'Behaviour'], [
    ['ANY', 'No filtering; detect edges of any polarity'],
    ['LIGHT_TO_DARK', 'Zero out positive gradient values (keep only bright→dark transitions)'],
    ['DARK_TO_LIGHT', 'Zero out negative gradient values (keep only dark→bright transitions)'],
])
doc.add_paragraph(
    'Direction compensation: For LEFT and TOP scan directions, the gradient kernel\'s intrinsic '
    'direction is opposite to the conceptual scan direction, so the polarity interpretation is '
    'flipped. This ensures "Light-to-Dark" always means "bright on the approach side" regardless '
    'of scan direction.')

doc.add_heading('4.6 Step 6: Edge Point Detection', level=2)
doc.add_paragraph('Functions: _detect_edge_point_horizontal / _detect_edge_point_vertical')
items = [
    'Apply border ignore: zero out gradient values within BORDER_IGNORE_PCT (default 2%) '
    'of image edges.',
    'Find indices where |gradient| > EDGE_THRESHOLD (default 25).',
    'Cluster Analysis: Identify contiguous clusters separated by gaps > MAX_CLUSTER_GAP (default 5 pixels).',
    'Cluster Selection: Based on scan direction — LEFT: first cluster from left; '
    'RIGHT: last cluster; TOP: first from top; BOTTOM: last from top.',
    'Find the peak within the selected cluster: argmax |gradient|.',
    'Sub-pixel refinement via 3-point parabolic fit:',
]
for item in items:
    add_numbered(item)
add_code('x_sub = x_peak + (g[x_peak-1] - g[x_peak+1]) / '
         '(2 × (g[x_peak-1] - 2×g[x_peak] + g[x_peak+1]))')
doc.add_paragraph('Offset clamped to [-0.5, +0.5]. The y-coordinate is the centre of the region strip.')

doc.add_heading('4.7 Step 7: RANSAC Line Fitting', level=2)
doc.add_paragraph('Function: fit_line_ransac(points, iterations, threshold)')
items = [
    'Requires ≥ 3 detected points.',
    'For RANSAC_ITERATIONS (default 2000) iterations: randomly sample 2 points, '
    'compute line Ax + By + C = 0, count inliers (perpendicular distance < RANSAC_THRESHOLD).',
    'Keep the model with the most inliers.',
    'Refine with cv2.fitLine (L2 distance) on the inlier set.',
    'Extract direction vector (vx, vy), point (x0, y0), and compute line endpoints at image boundaries.',
]
for item in items:
    add_numbered(item)

doc.add_heading('Line Equation', level=3)
add_code('A = vy,  B = -vx,  C = -vy × x0 + vx × y0')

doc.add_heading('4.8 Step 8: Perpendicular Delta Computation', level=2)
doc.add_paragraph('Function: compute_perpendicular_delta(A, B, C, img_h, img_w, scale)')
doc.add_paragraph(
    'Matches the C# EmguVision.cs implementation (WEPLEFT_REQ lines 776–787):')
items = [
    'Image centre in downsampled space: (cx, cy) = (W/2, H/2).',
    'Foot of perpendicular from centre to line:',
]
for item in items:
    add_numbered(item)
add_code('t = (A×cx + B×cy + C) / (A² + B²)\n'
         'px = cx - A×t,  py = cy - B×t')
items = [
    'Convert to original-resolution pixels: px_orig = px / s, py_orig = py / s.',
    'Delta: Δx = px_orig − cx_orig,  Δy = py_orig − cy_orig.',
]
for item in items:
    add_numbered(item)
doc.add_paragraph(
    'These deltas represent the offset of the wafer edge from the camera\'s optical centre, '
    'which the motion controller uses to align the wafer.')

# ===== 5. CONFIG REFERENCE =====
doc.add_heading('5. Configuration Reference', level=1)
add_table(['Parameter', 'Default', 'Description'], [
    ['SCAN_DIRECTION', 'LEFT', 'Edge side: LEFT, RIGHT, TOP, BOTTOM'],
    ['EDGE_POLARITY', 'ANY', 'Polarity filter: ANY, LIGHT_TO_DARK, DARK_TO_LIGHT'],
    ['NUM_REGIONS', '40', 'Number of scan strips'],
    ['EDGE_THRESHOLD', '25', 'Minimum gradient magnitude for edge'],
    ['MAX_CLUSTER_GAP', '5', 'Max gap between cluster points (px)'],
    ['BORDER_IGNORE_PCT', '0.02', 'Image border exclusion zone (2%)'],
    ['RANSAC_ITERATIONS', '2000', 'RANSAC iteration count'],
    ['RANSAC_THRESHOLD', '5.0', 'Inlier distance threshold (px)'],
    ['KERNEL_SIZE', '7', 'Gradient kernel width (odd, ≥ 3)'],
    ['TARGET_PROCESS_DIM', '1000', 'Resize target for processing'],
    ['CLAHE_CLIP_LIMIT', '2.0', 'CLAHE clip limit (0=off)'],
    ['CLAHE_GRID_SIZE', '8', 'CLAHE tile grid size'],
])

# ===== 6. DATA STRUCTURES =====
doc.add_heading('6. Data Structures', level=1)

doc.add_heading('6.1 EdgeFinderConfig', level=2)
doc.add_paragraph(
    'Extends ClassificationConfig. All edge-specific parameters listed in §5 plus inherited '
    'FOV classification settings.')

doc.add_heading('6.2 find_edge() Return Dict', level=2)
add_code("{\n"
         "  'success': bool,\n"
         "  'line_params': {'a':float,'b':float,'c':float,\n"
         "                  'vx':float,'vy':float,'x0':float,'y0':float},\n"
         "  'line_endpoints': {'x_top':int,'x_bot':int},  # or y_left,y_right\n"
         "  'slope': float, 'intercept_c': float,\n"
         "  'delta_x': float, 'delta_y': float,  # original-resolution offset\n"
         "  'intercept_point': {'x':float, 'y':float},\n"
         "  'detected_points': [(x,y), ...],\n"
         "  'inliers': np.array, 'num_points': int, 'num_inliers': int,\n"
         "  'region_data': [...], 'image': np.array,\n"
         "  'original_image': np.array, 'scale': float,\n"
         "  'scan_direction': str, 'is_vertical_edge': bool,\n"
         "  'reason': str or None\n"
         "}")

# ===== 7. ZMQ INTEGRATION =====
doc.add_heading('7. ZMQ Integration', level=1)

doc.add_heading('7.1 WAFER_EDGE_REQ Command', level=2)
add_code('WAFER_EDGE_REQ "<image_path>" "<recipe_path>" <direction> [<polarity>] [FORCE_RUN]')
doc.add_paragraph('Parameters:')
add_bullet('image_path: Full path to the image to analyse.')
add_bullet('recipe_path: Recipe XML (name or full path).')
add_bullet('direction: LEFT | RIGHT | TOP | BOTTOM.')
add_bullet('polarity: (optional) ANY | LIGHT_TO_DARK | DARK_TO_LIGHT.')
add_bullet('FORCE_RUN: (optional) Skip FOV classification.')

doc.add_paragraph('Response variants:')
add_code('{"status":"ok", "delta_x":-123.4, "delta_y":5.6,\n'
         ' "fov_type":"EDGE_FOV", "fov_confidence":0.95,\n'
         ' "a":..., "b":..., "c":..., "x_top":..., "x_bot":...}')
add_code('{"status":"die_fov_warning", "fov_type":"DIE_FOV",\n'
         ' "message":"Resend with FORCE_RUN to override."}')
add_code('{"status":"no_edge", "reason":"Not enough points"}')

doc.add_heading('7.2 DIE_FOV Warning Flow', level=2)
items = [
    'C# sends WAFER_EDGE_REQ without FORCE_RUN.',
    'If FOV classifier returns DIE_FOV, Python replies with die_fov_warning.',
    'C# shows a dialog to the operator.',
    'If operator confirms, C# resends with FORCE_RUN appended.',
    'Python runs edge detection regardless of FOV type.',
]
for item in items:
    add_numbered(item)

# ===== 8. RECIPE XML =====
doc.add_heading('8. Recipe XML Schema (FindWaferEdge)', level=1)
add_code('<FindWaferEdge>\n'
         '  <UseUniversalParams>false</UseUniversalParams>\n'
         '  <WaferAlignLeftParam Name="Wafer Align Left Parm">\n'
         '    <KernelSize>9</KernelSize>\n'
         '    <EdgeThreshold>500</EdgeThreshold>\n'
         '    <NumRegions>30</NumRegions>\n'
         '    <BorderIgnorePct>0.05</BorderIgnorePct>\n'
         '    <RansacThreshold>3.0</RansacThreshold>\n'
         '    <EdgePolarity>ANY</EdgePolarity>\n'
         '  </WaferAlignLeftParam>\n'
         '  <!-- Right, Top, Bottom follow same structure -->\n'
         '</FindWaferEdge>')
doc.add_paragraph(
    'Each scan direction has independent parameter sets. The UseUniversalParams flag (when true) '
    'applies LEFT parameters to all directions.')

# ===== 9. PER-DIRECTION CACHING =====
doc.add_heading('9. Per-Direction Config Caching', level=1)
doc.add_paragraph(
    'The AppState maintains a dict edge_configs with keys LEFT/RIGHT/TOP/BOTTOM. '
    'When the user switches the direction dropdown:')
items = [
    'EdgeViewModel.save_current_dir_to_cache(old_dir) saves current slider values.',
    'EdgeViewModel.get_cache_for_dir(new_dir) loads cached values for the new direction.',
    'The View updates slider positions without triggering detection.',
]
for item in items:
    add_numbered(item)

# ===== 10. FOV CLASSIFIER DETAILS =====
doc.add_heading('10. FOV Classifier Details', level=1)

doc.add_heading('10.1 Classification Config', level=2)
add_table(['Parameter', 'Default', 'Description'], [
    ['RELATIVE_CHANGE_WEAK', '0.15', '15% of dynamic range for weak edge'],
    ['RELATIVE_CHANGE_STRONG', '0.25', '25% for strong edge'],
    ['MIN_INTENSITY_CHANGE_WEAK', '30', 'Absolute fallback threshold'],
    ['MIN_INTENSITY_CHANGE_STRONG', '50', 'Absolute fallback threshold'],
    ['MIN_LENGTH_RATIO_LONG', '0.15', 'Minimum monotonic region length'],
    ['MIN_LENGTH_RATIO_SHORT', '0.05', 'Minimum for sharp edges'],
    ['SOBEL_STRENGTH_THRESHOLD', '0.20', '2-D Sobel edge threshold'],
    ['CONFIDENCE_THRESHOLD', '0.40', 'Below this → UNCERTAIN'],
    ['DIE_OVERRIDE_PCT', '100', '% regions that must be DIE to override'],
])

doc.add_heading('10.2 1-D Edge Detection Criteria', level=2)
doc.add_paragraph('Five criteria checked in order (first match wins):')
items = [
    'Curved Rising: Long monotonic rise (≥ 15% of profile), weak threshold, not periodic.',
    'Curved Falling: Same for falling.',
    'Sharp Rising: Short rise (≥ 5%), strong threshold, not periodic.',
    'Sharp Falling: Same for falling.',
    'Quarter Difference: First-quarter vs last-quarter mean intensity difference exceeds threshold.',
]
for item in items:
    add_numbered(item)

# ===== 11. C# COMPAT =====
doc.add_heading('11. C# Compatibility Notes', level=1)
add_bullet('Delta calculation matches EmguVision.cs lines 776–787 (perpendicular-intersection method).')
add_bullet('ZMQ protocol uses the same command format as the C# VPServer process commands.')
add_bullet('Recipe XML schema matches the C# WaferAlignLeftParam / WaferAlignRightParam structure.')
add_bullet('Direction-specific parameter storage mirrors C# per-side settings.')

# ===== 12. LIMITATIONS =====
doc.add_heading('12. Known Limitations & Future Work', level=1)
add_bullet('Curved wafer edges (large FOV) may produce slightly non-linear edge lines; '
           'a polynomial fit could improve accuracy.')
add_bullet('The FOV classifier occasionally misclassifies strong die patterns near the wafer edge.')
add_bullet('Gradient kernel is fixed; adaptive kernel selection based on image noise could improve robustness.')
add_bullet('No GPU acceleration (edge detection is fast enough on CPU at ~50 ms for 1000 px processed images).')

# ===== 13. FILE REFERENCE =====
doc.add_heading('13. File Reference', level=1)
add_table(['File', 'Purpose'], [
    ['app/services/edge_finder.py', 'Core edge detection algorithm (838 lines)'],
    ['app/services/fov_classifier.py', 'FOV classification (1418 lines)'],
    ['app/viewmodels/edge_viewmodel.py', 'UI business logic, config caching (208 lines)'],
    ['app/views/edge_tab.py', 'Tkinter tab UI'],
    ['app/models/recipe_model.py', 'XML recipe read/write'],
    ['app/models/app_state.py', 'Per-direction config cache'],
    ['app/services/zmq_server.py', 'ZMQ server (WAFER_EDGE_REQ handler)'],
])

# Save
out = os.path.join(os.path.dirname(__file__), 'Gradient_Based_Wafer_Edge_Detection.docx')
doc.save(out)
print(f'Saved: {out}')
